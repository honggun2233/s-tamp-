import sys
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ── 페이지 여백 ──────────────────────────────────────────
section = doc.sections[0]
section.page_width    = Cm(21)
section.page_height   = Cm(29.7)
section.left_margin   = Cm(2.5)
section.right_margin  = Cm(2.5)
section.top_margin    = Cm(2.5)
section.bottom_margin = Cm(2.5)

# ── 색상 ─────────────────────────────────────────────────
BLUE_DARK  = RGBColor(0x00, 0x41, 0x8A)
BLUE_LIGHT = RGBColor(0x1A, 0x73, 0xE8)
WHITE      = RGBColor(0xFF, 0xFF, 0xFF)
BLACK      = RGBColor(0x1A, 0x1A, 0x2E)
GRAY_TEXT  = RGBColor(0x55, 0x55, 0x55)

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def add_run(para, text, bold=False, size=10, color=None, italic=False):
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.name = '맑은 고딕'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    if color:
        run.font.color.rgb = color
    return run

def section_divider(title):
    doc.add_paragraph()
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = 'Table Grid'
    cell = tbl.cell(0, 0)
    set_cell_bg(cell, '004188')
    p2 = cell.paragraphs[0]
    p2.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p2.paragraph_format.space_before = Pt(4)
    p2.paragraph_format.space_after  = Pt(4)
    add_run(p2, f'  {title}', bold=True, size=11, color=WHITE)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

def scale_table(qnum, question_text, low_label, high_label, required=True):
    req = ' *' if required else ''
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(2)
    add_run(p, f'Q{qnum}.', bold=True, size=10, color=BLUE_DARK)
    add_run(p, f' {question_text}{req}', size=10, color=BLACK)

    tbl = doc.add_table(rows=3, cols=7)
    tbl.style = 'Table Grid'
    tbl.autofit = False
    col_widths = [Cm(4.0), Cm(1.8), Cm(1.8), Cm(1.8), Cm(1.8), Cm(1.8), Cm(4.0)]
    for i, w in enumerate(col_widths):
        for row in tbl.rows:
            row.cells[i].width = w

    # 행1: 레이블
    labels = ['', '매우\n불만족', '불만족', '보통', '만족', '매우\n만족', '']
    overrides = {0: low_label, 6: high_label}
    for i, cell in enumerate(tbl.rows[0].cells):
        set_cell_bg(cell, '004188' if i in [0, 6] else 'E8F0FE')
        p2 = cell.paragraphs[0]
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        txt = overrides.get(i, labels[i])
        clr = WHITE if i in [0, 6] else BLUE_DARK
        add_run(p2, txt, bold=True, size=9, color=clr)

    # 행2: 숫자
    for i, cell in enumerate(tbl.rows[1].cells):
        set_cell_bg(cell, 'F5F7FA')
        p2 = cell.paragraphs[0]
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if 1 <= i <= 5:
            add_run(p2, str(i), bold=True, size=11, color=BLUE_DARK)

    # 행3: 체크
    for i, cell in enumerate(tbl.rows[2].cells):
        set_cell_bg(cell, 'FFFFFF')
        p2 = cell.paragraphs[0]
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if 1 <= i <= 5:
            add_run(p2, '○', size=14, color=GRAY_TEXT)

    doc.add_paragraph().paragraph_format.space_after = Pt(2)

def choice_question(qnum, question_text, options, multi=False, required=True, note=''):
    req = ' *' if required else ''
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(2)
    add_run(p, f'Q{qnum}.', bold=True, size=10, color=BLUE_DARK)
    add_run(p, f' {question_text}{req}', size=10, color=BLACK)
    if note:
        add_run(p, f'  {note}', italic=True, size=9, color=GRAY_TEXT)

    sym = '☐' if multi else '○'
    for opt in options:
        op = doc.add_paragraph()
        op.paragraph_format.left_indent = Cm(0.8)
        op.paragraph_format.space_after = Pt(1)
        add_run(op, f'{sym} {opt}', size=10, color=GRAY_TEXT)

def open_question(qnum, question_text, lines=5, required=False):
    req = ' *' if required else ''
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(2)
    add_run(p, f'Q{qnum}.', bold=True, size=10, color=BLUE_DARK)
    add_run(p, f' {question_text}{req}', size=10, color=BLACK)

    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = 'Table Grid'
    cell = tbl.cell(0, 0)
    cell.width = Cm(16)
    set_cell_bg(cell, 'F5F7FA')
    cp = cell.paragraphs[0]
    for _ in range(lines):
        cp.add_run('\n')
    cp.paragraph_format.space_after = Pt(6)

# ════════════════════════════════════════════════════════
#  표지
# ════════════════════════════════════════════════════════
cover = doc.add_paragraph()
cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
cover.paragraph_format.space_before = Pt(60)
add_run(cover, 'S-TAMP 전자서명 시스템\n', bold=True, size=22, color=BLUE_DARK)
add_run(cover, '사용자 만족도 설문조사', bold=True, size=18, color=BLUE_LIGHT)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.paragraph_format.space_before = Pt(12)
add_run(sub, '시스템 통합 테스트(SI Test) | 삼성자산운용 디지털혁신팀', size=10, color=GRAY_TEXT)

date_p = doc.add_paragraph()
date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
date_p.paragraph_format.space_before = Pt(6)
add_run(date_p, '2026년 4월', size=10, color=GRAY_TEXT)

doc.add_page_break()

# ── 안내문 ────────────────────────────────────────────
tbl = doc.add_table(rows=1, cols=1)
tbl.style = 'Table Grid'
cell = tbl.cell(0, 0)
set_cell_bg(cell, 'E8F0FE')
cp = cell.paragraphs[0]
cp.paragraph_format.space_before = Pt(6)
cp.paragraph_format.space_after  = Pt(6)
cp.paragraph_format.left_indent  = Cm(0.5)
add_run(cp,
    '안녕하세요. 삼성자산운용 S-TAMP 전자서명 시스템 통합 테스트에 참여해 주셔서 감사합니다.\n'
    '본 설문은 시스템 사용 경험을 수집하여 서비스 품질 개선 및 성과 평가에 활용하기 위한 것입니다.\n'
    '응답 내용은 통계 목적으로만 활용되며, 소요 시간은 약 5분입니다.\n\n'
    '  * 표시 항목은 필수 응답입니다.',
    size=10, color=BLACK)
doc.add_paragraph().paragraph_format.space_after = Pt(6)

# ════════════════════════════════════════════════════════
#  섹션 1 — 기본 정보 (사용 빈도만)
# ════════════════════════════════════════════════════════
section_divider('섹션 1.  기본 정보')

choice_question(1, 'S-TAMP 시스템 사용 빈도는 어느 정도입니까?', [
    '하루 5회 이상',
    '하루 1~4회',
    '주 2~4회',
    '주 1회 이하',
])

# ════════════════════════════════════════════════════════
#  섹션 2 — 사용 편의성 (로그인·모바일 제외)
# ════════════════════════════════════════════════════════
section_divider('섹션 2.  시스템 사용 편의성')

scale_table(2,
    'S-TAMP 화면 구성(UI)이 직관적이고 이해하기 쉬웠습니까?',
    '매우 어려움', '매우 쉬움')

scale_table(3,
    '문서 템플릿 선택 및 내용 입력 과정이 편리하였습니까?',
    '매우 불편', '매우 편리')

scale_table(4,
    '기존 11단계 대비 3단계로 간소화된 새 프로세스(템플릿 선택 → 전자서명 → 자동발송)가 간결하게 느껴집니까?',
    '매우 복잡', '매우 간결')

# ════════════════════════════════════════════════════════
#  섹션 3 — Knox-ATIS 이중결재 통합
# ════════════════════════════════════════════════════════
section_divider('섹션 3.  Knox-ATIS 이중결재 통합 기능 평가')

scale_table(5,
    'Knox와 ATIS의 이중결재가 S-TAMP로 통합되어 One-Stop 처리가 가능해진 것에 만족하십니까?',
    '매우 불만족', '매우 만족')

scale_table(6,
    '기존 이중결재(Knox → ATIS)와 비교하여 업무 처리 시간이 단축되었다고 느끼십니까?',
    '전혀 못 느낌', '매우 크게 단축')

choice_question(7,
    'S-TAMP 통합 후 가장 크게 개선된 점은 무엇입니까?',
    ['중복 입력 해소', '결재 대기 시간 단축', '시스템 간 이동 불필요', '휴먼 에러 감소', '기타: ____________'],
    multi=True, note='(복수 선택 가능)')

scale_table(8,
    '결재 워크플로(상신 → 승인 → 완료) 흐름이 명확하고 이해하기 쉽습니까?',
    '전혀 명확하지 않음', '매우 명확함')

# ════════════════════════════════════════════════════════
#  섹션 4 — 시스템 성능 및 안정성
# ════════════════════════════════════════════════════════
section_divider('섹션 4.  시스템 성능 및 안정성')

scale_table(9,
    '전자서명 처리 속도(서명 요청 ~ 완료 소요 시간)에 만족하십니까?',
    '매우 불만족', '매우 만족')

choice_question(10,
    '테스트 기간 중 시스템 오류 또는 장애를 경험하셨습니까?', [
        '없었다',
        '1~2회 경험하였다',
        '3~5회 경험하였다',
        '5회 이상 자주 경험하였다',
    ])

choice_question(11,
    '(오류 경험 시) 오류 안내 메시지가 원인과 해결 방법을 충분히 안내하였습니까?', [
        '해당 없음 (오류 없었음)',
        '전혀 충분하지 않았다',
        '다소 부족하였다',
        '보통이었다',
        '충분하였다',
        '매우 충분하였다',
    ], required=False)

# ════════════════════════════════════════════════════════
#  섹션 5 — 보안 및 컴플라이언스
# ════════════════════════════════════════════════════════
section_divider('섹션 5.  보안 및 컴플라이언스')

scale_table(12,
    '블록체인 기반 서명 이력 관리(Audit Trail)가 인감 오남용 방지에 효과적이라고 생각하십니까?',
    '전혀 효과없음', '매우 효과적')

scale_table(13,
    'S-TAMP의 전자서명이 법적 효력을 갖추고 있다는 것에 신뢰가 가십니까?',
    '전혀 신뢰 안 함', '매우 신뢰함')

scale_table(14,
    '서명 완료 문서의 보관·조회 기능(중앙 저장소)이 충분합니까?',
    '매우 부족', '매우 충분')

# ════════════════════════════════════════════════════════
#  섹션 6 — 업무 개선 효과
# ════════════════════════════════════════════════════════
section_divider('섹션 6.  업무 개선 효과 평가')

scale_table(15,
    '기존 종이 인감 프로세스 대비 S-TAMP 도입으로 전체 업무 효율이 향상되었다고 느끼십니까?',
    '전혀 향상 안 됨', '매우 크게 향상')

scale_table(16,
    '출력 → 날인 → 스캔 → 우편 발송 단계가 사라진 것이 실무에서 체감상 도움이 됩니까?',
    '전혀 도움 안 됨', '매우 도움됨')

scale_table(17,
    '고객 문서 처리 속도(기존 3일 이상 → 당일 처리)에 만족하십니까?',
    '매우 불만족', '매우 만족')

# ════════════════════════════════════════════════════════
#  섹션 7 — 종합 만족도
# ════════════════════════════════════════════════════════
section_divider('섹션 7.  종합 만족도 및 의견')

scale_table(18,
    'S-TAMP 전자서명 시스템에 대한 전반적인 만족도를 평가해 주십시오.',
    '매우 불만족', '매우 만족')

scale_table(19,
    '동료/타 부서에 S-TAMP 사용을 권장하시겠습니까? (NPS)',
    '절대 권장 안 함', '적극 권장함')

choice_question(20,
    '가장 만족스러운 기능은 무엇입니까?',
    ['Knox-ATIS 이중결재 통합', '3단계 간소화 프로세스',
     '블록체인 Audit Trail', '당일 즉시 처리',
     '중앙 문서 보관·조회', '기타: ____________'],
    multi=True, note='(복수 선택 가능)')

choice_question(21,
    '가장 시급하게 개선이 필요한 부분은 무엇입니까?',
    ['화면 UI 구성', '처리 속도', '오류 안내 메시지',
     '문서 검색·조회', '교육/도움말', '기타: ____________'],
    multi=True, note='(복수 선택 가능)')

open_question(22,
    'S-TAMP 시스템 개선을 위한 자유로운 의견을 남겨 주세요.',
    lines=5)

# ── 마무리 ───────────────────────────────────────────
doc.add_paragraph()
thanks = doc.add_paragraph()
thanks.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(thanks,
    '소중한 의견을 주셔서 감사합니다.\n'
    '귀하의 응답은 S-TAMP 시스템 개선 및 전사 확대 적용에 적극 반영하겠습니다.',
    size=10, color=GRAY_TEXT)

footer_p = doc.add_paragraph()
footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
footer_p.paragraph_format.space_before = Pt(8)
add_run(footer_p, '삼성자산운용 디지털혁신팀  |  2026. 04', size=9, color=GRAY_TEXT)

# ── 저장 ─────────────────────────────────────────────
out = r'C:\project\전자서명 프로젝트\S-TAMP_사용자만족도_설문조사.docx'
doc.save(out)
print(f'저장 완료: {out}')
